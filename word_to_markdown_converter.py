#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word转Markdown转换器
专门用于将Word文档转换为Markdown格式，保持原有格式和结构
支持标题、段落、表格、列表等元素的准确转换
"""

import os
import re
from typing import List, Dict, Any, Optional
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


class WordToMarkdownConverter:
    """Word文档转Markdown转换器"""
    
    def __init__(self):
        self.markdown_content = []
        self.current_list_level = 0
        
    def convert_document(self, docx_path: str, output_path: str = None) -> str:
        """转换Word文档为Markdown格式
        
        Args:
            docx_path: Word文档路径
            output_path: 输出Markdown文件路径（可选）
            
        Returns:
            转换后的Markdown内容
        """
        try:
            # 读取Word文档
            doc = Document(docx_path)
            print(f"正在转换文档: {docx_path}")
            
            # 重置内容
            self.markdown_content = []
            
            # 处理文档元素
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # 处理段落
                    paragraph = self._find_paragraph_by_element(doc, element)
                    if paragraph:
                        self._process_paragraph(paragraph)
                elif element.tag.endswith('tbl'):
                    # 处理表格
                    table = self._find_table_by_element(doc, element)
                    if table:
                        self._process_table(table)
            
            # 生成最终的Markdown内容
            markdown_text = self._finalize_markdown()
            
            # 保存到文件
            if output_path:
                self._save_markdown(markdown_text, output_path)
                print(f"转换完成，已保存到: {output_path}")
            
            return markdown_text
            
        except Exception as e:
            print(f"转换过程中出现错误: {str(e)}")
            raise
    
    def _find_paragraph_by_element(self, doc: DocumentType, element) -> Optional[Paragraph]:
        """根据XML元素查找对应的段落对象"""
        for paragraph in doc.paragraphs:
            if paragraph._element == element:
                return paragraph
        return None
    
    def _find_table_by_element(self, doc: DocumentType, element) -> Optional[Table]:
        """根据XML元素查找对应的表格对象"""
        for table in doc.tables:
            if table._element == element:
                return table
        return None
    
    def _process_paragraph(self, paragraph: Paragraph):
        """处理段落元素"""
        text = paragraph.text.strip()
        if not text:
            # 空段落作为换行
            self.markdown_content.append("")
            return
        
        # 检测标题级别
        heading_level = self._get_heading_level(paragraph)
        if heading_level > 0:
            # 处理标题
            self.markdown_content.append(f"{'#' * heading_level} {text}")
            self.markdown_content.append("")
        else:
            # 检测列表
            if self._is_list_item(paragraph):
                self._process_list_item(paragraph)
            else:
                # 普通段落
                self.markdown_content.append(text)
                self.markdown_content.append("")
    
    def _get_heading_level(self, paragraph: Paragraph) -> int:
        """获取段落的标题级别"""
        style_name = paragraph.style.name.lower()
        
        # 检查样式名称
        if 'heading' in style_name:
            # 提取数字
            match = re.search(r'(\d+)', style_name)
            if match:
                return int(match.group(1))
        
        # 检查字体大小和加粗
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold and first_run.font.size:
                size_pt = first_run.font.size.pt
                if size_pt >= 18:
                    return 1
                elif size_pt >= 16:
                    return 2
                elif size_pt >= 14:
                    return 3
        
        return 0
    
    def _is_list_item(self, paragraph: Paragraph) -> bool:
        """检测是否为列表项"""
        text = paragraph.text.strip()
        # 检查是否以数字、字母或符号开头
        list_patterns = [
            r'^\d+[.)]\s+',  # 数字列表
            r'^[a-zA-Z][.)]\s+',  # 字母列表
            r'^[•·▪▫-]\s+',  # 符号列表
            r'^\([\d\w]+\)\s+',  # 括号列表
        ]
        
        for pattern in list_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _process_list_item(self, paragraph: Paragraph):
        """处理列表项"""
        text = paragraph.text.strip()
        
        # 检测列表类型和级别
        if re.match(r'^\d+[.)]\s+', text):
            # 有序列表
            content = re.sub(r'^\d+[.)]\s+', '', text)
            self.markdown_content.append(f"1. {content}")
        else:
            # 无序列表
            content = re.sub(r'^[•·▪▫-]\s+', '', text)
            content = re.sub(r'^\([\d\w]+\)\s+', '', content)
            content = re.sub(r'^[a-zA-Z][.)]\s+', '', content)
            self.markdown_content.append(f"- {content}")
    
    def _process_table(self, table: Table):
        """处理表格元素"""
        if not table.rows:
            return
        
        # 添加表格前的空行
        self.markdown_content.append("")
        
        # 处理表格行
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                # 清理单元格文本
                cell_text = self._clean_cell_text(cell.text)
                row_cells.append(cell_text)
            
            # 生成Markdown表格行
            markdown_row = "| " + " | ".join(row_cells) + " |"
            self.markdown_content.append(markdown_row)
            
            # 在第一行后添加分隔符
            if row_idx == 0:
                separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                self.markdown_content.append(separator)
        
        # 添加表格后的空行
        self.markdown_content.append("")
    
    def _clean_cell_text(self, text: str) -> str:
        """清理单元格文本"""
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text.strip())
        # 转义Markdown特殊字符
        text = text.replace('|', '\\|')
        text = text.replace('\n', ' ')
        return text
    
    def _finalize_markdown(self) -> str:
        """生成最终的Markdown内容"""
        # 清理多余的空行
        cleaned_content = []
        prev_empty = False
        
        for line in self.markdown_content:
            if line.strip() == "":
                if not prev_empty:
                    cleaned_content.append("")
                prev_empty = True
            else:
                cleaned_content.append(line)
                prev_empty = False
        
        # 移除开头和结尾的空行
        while cleaned_content and cleaned_content[0] == "":
            cleaned_content.pop(0)
        while cleaned_content and cleaned_content[-1] == "":
            cleaned_content.pop()
        
        return "\n".join(cleaned_content)
    
    def _save_markdown(self, content: str, output_path: str):
        """保存Markdown内容到文件"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        except Exception as e:
            print(f"保存文件时出现错误: {str(e)}")
            raise


def main():
    """主函数"""
    # 配置文件路径
    input_file = "银行一表通监管数据采集接口标准（2.0正式版）.docx"
    output_file = "银行一表通监管数据采集接口标准_2.0_word转换版.md"
    
    # 检查输入文件是否存在
    if not os.path.exists(input_file):
        print(f"错误: 找不到输入文件 {input_file}")
        return
    
    try:
        # 创建转换器并执行转换
        converter = WordToMarkdownConverter()
        result = converter.convert_document(input_file, output_file)
        
        print(f"\n转换成功完成!")
        print(f"输入文件: {input_file}")
        print(f"输出文件: {output_file}")
        print(f"生成内容长度: {len(result)} 字符")
        
    except Exception as e:
        print(f"转换失败: {str(e)}")


if __name__ == "__main__":
    main()